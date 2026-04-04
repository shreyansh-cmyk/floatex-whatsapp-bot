"""Document templates for Floatex Solar — generates actual .docx files."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import io


# --- Styling helpers ---

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color,
    })
    shading.append(shading_elm)


def add_header_footer(doc, project_title, doc_no):
    """Add header and footer to document."""
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Header
    header = section.header
    hp = header.paragraphs[0]
    hp.text = f"{project_title}\n{doc_no}"
    hp.style.font.size = Pt(8)
    hp.style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def styled_para(doc, text, size=11, bold=False, color=None, align=None, space_after=6):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Arial'
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def section_heading(doc, number, title):
    """Add a numbered section heading."""
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {title}")
    run.font.size = Pt(13)
    run.font.name = 'Arial'
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    return p


def bullet_point(doc, text, indent=0.5):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.text = text
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    return p


def add_table(doc, headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Arial'
        set_cell_shading(cell, 'D6E4F0')

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'

    return table


# --- Float Storage SOP ---

def generate_float_storage_sop(project):
    """Generate Float Storage SOP .docx for a project."""
    doc = Document()
    pid = project.get("id", "PXXX")
    name = project.get("name", "Project")
    full_name = project.get("full_name", name)
    capacity = project.get("capacity_mw", "XX")
    epc = project.get("epc", "EPC Contractor")
    site = project.get("site_location") or project.get("reservoir_name") or "Site Location"
    client = project.get("client_name") or epc
    doc_no = f"FSR-{pid}-SOP-SM-01_R0"

    project_title = f"{capacity} MW FLOATING SOLAR PV PROJECT\n{full_name}"

    add_header_footer(doc, project_title, doc_no)

    # --- Cover Page ---
    doc.add_paragraph()
    doc.add_paragraph()
    styled_para(doc, project_title, size=16, bold=True, color=(0x1F, 0x4E, 0x79), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    styled_para(doc, "SOP for Float Storage", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    styled_para(doc, doc_no, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    # Revision table
    add_table(doc,
        ["Rev", "Date", "Written By", "Checked By", "Approved By", "Description"],
        [["R0", "—", "—", "—", "—", "Issued for Information"]]
    )

    doc.add_paragraph()
    doc.add_paragraph()
    styled_para(doc, "SOP FOR FLOAT STORAGE", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # --- Contents ---
    styled_para(doc, "Contents", size=14, bold=True, color=(0x1F, 0x4E, 0x79))
    for i, title in enumerate([
        "INTRODUCTION", "SCOPE", "STORAGE LOCATION",
        "STORAGE METHODOLOGY", "STORAGE PROCEDURE", "STORAGE AREA SAFETY"
    ], 1):
        styled_para(doc, f"{i}. {title}", size=11)

    doc.add_page_break()

    # --- Section 1: Introduction (PROJECT-SPECIFIC) ---
    section_heading(doc, 1, "INTRODUCTION")

    styled_para(doc, (
        f"This Standard Operating Procedure (SOP) outlines the methodology and procedure "
        f"for storage of HDPE floaters at the project site for the {full_name} project."
    ))

    styled_para(doc, "Project Particulars", size=11, bold=True, space_after=4)
    add_table(doc,
        ["Particular", "Description"],
        [
            ["Project Name", full_name],
            ["Plant AC Capacity", f"{capacity} MW"],
            ["Project ID", pid],
            ["EPC Contractor", epc],
            ["Site Location", site],
            ["Client / Owner", client],
            ["Design Life", "25 years"],
            ["Floatex Scope", "Float system design, supply & supervision"],
        ]
    )

    styled_para(doc, (
        f"\nThrough the contractual process, the development of the {capacity} MW floating solar "
        f"PV project has been awarded to {epc}. {epc} has subcontracted the floating system "
        f"scope of work to Floatex Solar Pvt. Ltd."
    ), space_after=12)

    # --- Section 2: Scope (STANDARD) ---
    section_heading(doc, 2, "SCOPE")
    styled_para(doc, (
        "This document briefs the methodology and the procedure of storage of floaters "
        "being used for the project, as mentioned in Section 1."
    ))

    # --- Section 3: Storage Location (STANDARD) ---
    section_heading(doc, 3, "STORAGE LOCATION")
    styled_para(doc, (
        "Approximate 1 hectare area is required for storage of 40 MWp of floaters for the project. "
        "It is recommended to develop at least 50 MWp of storage area at two different locations "
        "near to the launch pad."
    ))
    styled_para(doc, (
        f"Development of storage area and location is in the scope of {epc}."
    ))
    styled_para(doc, (
        "The storage area layout should accommodate the following zones:"
    ))
    bullet_point(doc, "Float unloading zone (vehicle access)")
    bullet_point(doc, "Stacking zone (organized by float type)")
    bullet_point(doc, "Buffer zone (minimum 2m between stacks and boundary)")
    bullet_point(doc, "Access pathways for material handling equipment")

    # --- Section 4: Storage Methodology (STANDARD) ---
    section_heading(doc, 4, "STORAGE METHODOLOGY")
    styled_para(doc, (
        "The following methodology is adopted while storing the floats in the designated area:"
    ))
    bullet_point(doc, "The designated area must be cleaned and levelled prior to float storage. "
                      "The ground surface should be free from any sharp object, stone etc. which can damage the float.")
    bullet_point(doc, "Floats must be placed in stacked manner")
    bullet_point(doc, "Four numbers of floats must be strapped together to make one stack")
    bullet_point(doc, "Four such stacks will be kept on top of each other")
    bullet_point(doc, "The height of each stack should ensure that the floats in stacks are stable")
    bullet_point(doc, "Floats of the same type must be stored together — do not mix Panel Floats, "
                      "Aisle Floats, and Side Floats in the same stack")

    # --- Section 5: Storage Procedure (STANDARD) ---
    section_heading(doc, 5, "STORAGE PROCEDURE")
    styled_para(doc, (
        "The following procedure must be ensured for proper handling and storage "
        "of the floaters in the storage area:"
    ))
    bullet_point(doc, "Proper care should be given while unloading the floats from the vehicle "
                      "to ensure no damage occurs")
    bullet_point(doc, "Floats must NOT be thrown from the vehicle")
    bullet_point(doc, "Floats must NOT be dragged to the final stack location — instead should be "
                      "lifted and placed on the stack")
    bullet_point(doc, "Inspect each float visually during unloading for cracks, deformation, or damage")
    bullet_point(doc, "Damaged floats must be separated and reported immediately")
    bullet_point(doc, "Maintain a count register for each delivery — verify quantity against challan")

    # --- Section 6: Storage Area Safety (STANDARD) ---
    section_heading(doc, 6, "STORAGE AREA SAFETY")
    styled_para(doc, (
        "The following measures must be ensured for the safety of the storage area:"
    ))
    bullet_point(doc, "The designated area must be fenced from all sides except the main entry")
    bullet_point(doc, "Security must be deployed for the protection of the floaters")
    bullet_point(doc, "No smoking or open flames within 50m of the storage area (HDPE is flammable)")
    bullet_point(doc, "Fire extinguishers must be placed at entry and exit points")
    bullet_point(doc, "Adequate lighting for night-time security")
    bullet_point(doc, "Drainage must be ensured — waterlogging can displace stacks")

    # --- Footer ---
    doc.add_paragraph()
    styled_para(doc, "— End of Document —", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x99, 0x99, 0x99))
    styled_para(doc, f"FLOATEX DOC. NO.: {doc_no}", size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x99, 0x99, 0x99))

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, doc_no


# --- HIRA ---

# Standard JSA data — same for every FSPV project
JSA_DATA = [
    {
        "activity": "Site Mobilization / Vegetation Clearance",
        "hazards": [
            ("Snake bite, insect bites, honey bee stings", "Injuries, fatality", "Check for valid permit, organize local person, field first aid box with snake antidote, work during day hours only", "Gumboots, helmet, goggles, vest, gloves, nose mask"),
            ("Falling branches, trees", "Serious injuries, fatality", "Briefing before entry, identify line markings, exercise caution", "Safety helmet, full body harness"),
            ("Trip/fall on uneven terrain", "Injuries", "Clear pathways, adequate lighting, buddy system", "Safety shoes, helmet"),
        ],
    },
    {
        "activity": "Structural Materials & Equipment Shifting/Lifting",
        "hazards": [
            ("Pinching of hands/legs while handling", "Serious injuries", "Keep body parts away from pinch points, proper stacking", "Safety shoes, helmet, gloves"),
            ("Swinging of suspended load", "Serious injuries, fatality", "Use tag lines at both ends, qualified rigger/signalman", "Safety shoes, helmet, vest"),
            ("Overloading of crane/hydra", "Property damage, fatality", "Use within rated capacity, display SWL on all equipment", "Helmet, vest"),
            ("Failure of lifting tools/tackles", "Serious injuries", "Use certified/tested tools, daily inspection checklist", "Helmet, gloves, shoes"),
        ],
    },
    {
        "activity": "Erection of Structural Materials (Crane & Rigging)",
        "hazards": [
            ("Hit by load due to swinging", "Serious injuries, fatality", "Cordon off lifting area, no unauthorized entry, use tag lines", "Helmet, shoes, harness"),
            ("Pinching points during assembly", "Cut injuries", "Communication with coworker, proper supervision", "Gloves, helmet, shoes"),
            ("Dropped objects from height", "Injuries, property damage", "Tool belt for containment, barricade work area", "Helmet, shoes, goggles"),
        ],
    },
    {
        "activity": "Scaffolding Erection/Dismantling",
        "hazards": [
            ("Collapse of scaffold structure", "Serious injuries, fatality", "Only certified personnel, examine all components before use", "Helmet, harness, shoes"),
            ("Slips, trips and falls from scaffold", "Injuries", "Secure all platforms/planks, maintain safe distance from power lines", "Harness, helmet, shoes"),
            ("Dropped objects", "Injuries", "Use tool belt, secure tubes/planks when lifting/lowering, barricade area", "Helmet, shoes, goggles"),
        ],
    },
    {
        "activity": "Gas Cutting & Welding Work",
        "hazards": [
            ("Fire and explosion", "Burn injuries, property damage", "Flash back arrester on cylinder, remove combustibles, fire extinguisher ready", "Welding shield, goggles, gloves, apron"),
            ("Welding fumes inhalation", "Health hazard", "Adequate ventilation, nose mask, hot work permit", "Nose mask, goggles"),
            ("Gas cylinder fall/leak", "Explosion, injuries", "Check leakage with soap solution, use trolley for shifting, chain cylinders", "Helmet, shoes, gloves"),
            ("Eye injury from arc/spatters", "Eye damage", "Appropriate welding goggles for all nearby workers", "Welding goggles, face shield"),
        ],
    },
    {
        "activity": "Electrical Work",
        "hazards": [
            ("Electric shock", "Serious injury, fatality", "LOTO procedure, use insulated tools, qualified electrician only", "Insulated gloves, shoes, goggles"),
            ("Arc flash", "Burn injuries", "Maintain safe distance, use PPE rated for arc flash", "Arc-rated PPE, face shield"),
            ("Cable damage during laying", "Electric shock, fire", "Inspect cables before use, proper cable routing", "Gloves, shoes, helmet"),
        ],
    },
    {
        "activity": "Excavation Work (JCB/Manual)",
        "hazards": [
            ("Fall of person into excavation", "Serious injuries", "Install cordon/warning tape, no person during JCB operation", "Helmet, shoes, vest"),
            ("Soil collapse", "Burial, fatality", "Shoring for deep excavations, excavation permit required", "Helmet, shoes, gloves"),
            ("Underground utility strike", "Electric shock, gas leak", "Work area clearance certificate before starting", "Insulated shoes, gloves"),
        ],
    },
    {
        "activity": "Civil Work (Concrete/Casting)",
        "hazards": [
            ("Slippery approach/surfaces", "Slip/trip/fall", "Anti-slip footwear, clean walkways, warning signs", "Gumboots, helmet, gloves"),
            ("Chemical exposure from cement", "Skin irritation, respiratory issues", "Use gloves and mask when handling cement, wash exposed skin", "Gloves, nose mask, goggles"),
            ("Fall from height during formwork", "Serious injuries", "Full body harness, life line at height >2m", "Harness, helmet, shoes"),
            ("Sharp edges of construction materials", "Cut injuries", "Handle with care, use gloves, stack neatly", "Gloves, shoes"),
        ],
    },
    {
        "activity": "Anchor Block Casting & Dropping",
        "hazards": [
            ("Block fall during lifting/shifting", "Crush injuries, fatality", "Certified crane, rated slings, no person under load", "Helmet, shoes, vest"),
            ("Barge instability during loading", "Drowning, injuries", "Check barge stability, load as per plan, life jackets mandatory", "Life jacket, helmet, shoes"),
            ("DGPS equipment malfunction during dropping", "Incorrect anchor placement", "Calibrate DGPS daily, backup positioning method", "Life jacket, helmet"),
            ("Diver hazard during underwater verification", "Drowning", "Expert diver only, monitoring during dive, emergency rescue team ready", "Diving equipment, life jacket"),
        ],
    },
    {
        "activity": "Work on Water Body (Float Assembly, Towing, Mooring)",
        "hazards": [
            ("Fall into water from floats", "Drowning, fatality", "Life jackets mandatory, rescue boat with spare jackets always on standby", "Life jacket, helmet, shoes"),
            ("Trip/slip on wet float surface", "Injuries", "Anti-slip footwear, clear walkways on floats, no running", "Anti-slip shoes, life jacket"),
            ("Electrical hazard on water", "Electrocution", "All electrical connections above water level, GFCI protection", "Insulated gloves, life jacket"),
            ("Heat stress/sunburn", "Heat stroke", "Shade breaks every 2 hours, hydration stations, early morning start", "Hat, sunscreen, vest"),
            ("Mooring rope snap-back", "Serious injuries, fatality", "Stand clear of rope line, use rated ropes, inspect before each use", "Helmet, gloves, life jacket"),
        ],
    },
]


def generate_hira(project):
    """Generate HIRA .docx for a project."""
    doc = Document()
    pid = project.get("id", "PXXX")
    name = project.get("name", "Project")
    full_name = project.get("full_name", name)
    capacity = project.get("capacity_mw") or "XX"
    epc = project.get("epc", "EPC Contractor")
    site = project.get("site_location") or project.get("reservoir_name") or "Site"
    doc_no = f"FSR-{pid}-GEN-JS-01_R0"

    project_title = f"{capacity} MW FLOATING SOLAR PV PROJECT\n{full_name}"

    add_header_footer(doc, project_title, doc_no)

    # --- Cover Page ---
    doc.add_paragraph()
    doc.add_paragraph()
    styled_para(doc, project_title, size=16, bold=True, color=(0x1F, 0x4E, 0x79), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    styled_para(doc, "HAZARD IDENTIFICATION AND RISK ASSESSMENT\nfor FPV SYSTEM", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    styled_para(doc, doc_no, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    add_table(doc,
        ["Rev", "Date", "Written By", "Checked By", "Approved By", "Description"],
        [["R0", "—", "—", "—", "—", "Issued for Review"]]
    )

    doc.add_paragraph()
    styled_para(doc, "HAZARD IDENTIFICATION AND RISK ASSESSMENT for FPV SYSTEM", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # --- Section 1: Introduction (PROJECT-SPECIFIC) ---
    section_heading(doc, 1, "INTRODUCTION")
    styled_para(doc, (
        f"This document provides the Hazard Identification and Risk Assessment (HIRA) "
        f"for the floating photovoltaic (FPV) system activities at the {full_name} project."
    ))

    styled_para(doc, "Project Particulars", size=11, bold=True, space_after=4)
    add_table(doc,
        ["Particular", "Description"],
        [
            ["Project Name", full_name],
            ["Plant AC Capacity", f"{capacity} MW"],
            ["Project ID", pid],
            ["EPC Contractor", epc],
            ["Site Location", site],
            ["Design Life", "25 years"],
        ]
    )

    styled_para(doc, (
        f"\n{epc} has subcontracted the floating system scope to Floatex Solar Pvt. Ltd. "
        f"This HIRA covers all FPV system activities under Floatex's scope."
    ))

    # --- Section 2: Scope ---
    section_heading(doc, 2, "SCOPE")
    styled_para(doc, "This document provides guidance on Job Safety Analysis (JSA) for the FPV system, covering all major site activities from mobilization through commissioning.")

    # --- Section 3: JSA Tables ---
    section_heading(doc, 3, "JSA FOR FPV SYSTEM")

    for jsa in JSA_DATA:
        styled_para(doc, f"JOB Task: {jsa['activity']}", size=11, bold=True, color=(0x1F, 0x4E, 0x79), space_after=4)

        headers = ["Potential Hazards", "Risk/Hazard Effect", "Safety Precautions", "PPE"]
        rows = []
        for h in jsa["hazards"]:
            rows.append([h[0], h[1], h[2], h[3]])

        add_table(doc, headers, rows)
        doc.add_paragraph()

    # --- Section 4: General Safety Requirements ---
    section_heading(doc, 4, "GENERAL SAFETY REQUIREMENTS")
    for req in [
        "All workers must undergo safety induction before starting work on site",
        "Daily toolbox talk (TBT) must be conducted before each work shift",
        "Valid medical fitness certificates required for all workers",
        "Group insurance policy must cover all workers on site",
        "First aid kit must be available at every work location",
        "Emergency contact numbers displayed at prominent locations",
        "Incident reporting: all near-misses and incidents must be reported within 4 hours",
        "PPE compliance: 100% PPE usage enforced at all times within site boundary",
        "No work permitted during thunderstorm, heavy rain, or wind speed >50 km/h",
        "Rescue boat with trained swimmers must be on standby during all water body operations",
    ]:
        bullet_point(doc, req)

    # --- Footer ---
    doc.add_paragraph()
    styled_para(doc, "— End of Document —", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x99, 0x99, 0x99))
    styled_para(doc, f"FLOATEX DOC. NO.: {doc_no}", size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x99, 0x99, 0x99))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, doc_no


# --- Template registry ---

TEMPLATES = {
    "float-storage-sop": {
        "title": "Float Storage SOP",
        "generator": generate_float_storage_sop,
        "filename_pattern": "FSR-{pid}-SOP-SM-01_R0.docx",
    },
    "hira": {
        "title": "HIRA — Hazard Identification & Risk Assessment",
        "generator": generate_hira,
        "filename_pattern": "FSR-{pid}-GEN-JS-01_R0.docx",
    },
}
